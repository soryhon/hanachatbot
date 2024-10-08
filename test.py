import streamlit as st
import pandas as pd
import os
import openai
from langchain.llms import OpenAI
from PIL import Image
from PyPDF2 import PdfReader
from io import BytesIO
import docx
import xlsxwriter

# OpenAI API 설정 (올바른 API 키로 설정하세요)
os.environ["OPENAI_API_KEY"] = "sk-iQH3k8Kwr4aizl_6HEldJsCmVkDMWiMUnNE2eGe6KsT3BlbkFJQ8YYozec2S8tpnzNDWKHb3S8xVbqw63fUTNTh9wAoA"
llm = OpenAI(model='gpt-4', temperature=0.7)

# 메인 레이아웃
st.title("일일 업무 및 보고서 자동화 프로그램")

# 세 개의 프레임을 나누어서 배치
col1, col2, col3 = st.columns([4, 1, 5])

# 1. 작성 보고서 요청사항 (col1에 위치)
with col1:
    st.header("1. 작성 보고서 요청사항")
    
    # pandas 데이터프레임을 사용하여 입력받을 데이터 구조 생성
    if 'df' not in st.session_state:
        st.session_state['df'] = pd.DataFrame(columns=['제목', '요청', '데이터'])

    # 파일 선택 함수
    def select_file(row_idx):
        uploaded_file = st.file_uploader(f"데이터 선택 (행 {row_idx})", type=['csv', 'txt', 'pdf', 'docx', 'pptx', 'xlsx'], key=f'file_{row_idx}')
        if uploaded_file is not None:
            file_path = os.path.join("/mnt/data", uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.session_state['df'].at[row_idx, '데이터'] = file_path
            st.success(f"파일 저장 완료: {file_path}")

    # 데이터프레임의 각 행을 입력받는 인터페이스 생성
    for idx, row in st.session_state['df'].iterrows():
        col_title, col_request, col_data = st.columns([3, 3, 4])
        col_title.text_input("제목", key=f'title_{idx}', value=row['제목'] if pd.notna(row['제목']) else "")
        col_request.text_input("요청", key=f'request_{idx}', value=row['요청'] if pd.notna(row['요청']) else "")
        col_data.button("데이터 선택", on_click=lambda idx=idx: select_file(idx))

    # 행 추가 함수
    def add_row():
        new_row = {'제목': '', '요청': '', '데이터': ''}
        # DataFrame에 새로운 행 추가
        st.session_state['df'] = pd.concat([st.session_state['df'], pd.DataFrame([new_row])], ignore_index=True)

    # 행 추가 및 삭제 버튼
    st.button("행 추가", on_click=add_row)
    selected_rows = st.multiselect("삭제할 행 선택", st.session_state['df'].index)

    # 행 삭제 함수
    def delete_selected_rows():
        st.session_state['df'] = st.session_state['df'].drop(selected_rows).reset_index(drop=True)

    st.button("행 삭제", on_click=delete_selected_rows)

# 2. 파일 업로드 (col1에 위치)
with col1:
    st.header("2. 파일 업로드")
    uploaded_files = st.file_uploader("파일 선택", accept_multiple_files=True)
    if uploaded_files:
        for file in uploaded_files:
            st.write(f"업로드된 파일명: {file.name}")
            # 파일 저장 로직 필요

# 3. 실행 (col2에 위치)
with col2:
    st.header("3. 실행")
    if st.button("실행"):
        if 'df' in st.session_state:
            results = []
            for idx, row in st.session_state['df'].iterrows():
                prompt = f"제목: {row['제목']}, 요청: {row['요청']}, 데이터: {row['데이터']}"
                response = llm(prompt)
                results.append(response)
            st.session_state['results'] = results
            st.success("LLM 처리 완료")

# 4. 결과 보고서 (col3에 위치)
with col3:
    st.header("4. 결과 보고서")
    if 'results' in st.session_state:
        for idx, result in enumerate(st.session_state['results']):
            st.subheader(f"제목({st.session_state['df'].at[idx, '제목']})")
            st.write(result)

    # 파일 형식 선택
    export_format = st.selectbox("저장할 형식 선택", ['pdf', 'docx', 'xlsx', 'txt'])

    if st.button("Export"):
        if 'results' in st.session_state:
            # PDF 저장
            if export_format == 'pdf':
                from fpdf import FPDF
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                for idx, result in enumerate(st.session_state['results']):
                    pdf.cell(200, 10, txt=f"제목: {st.session_state['df'].at[idx, '제목']}", ln=True)
                    pdf.multi_cell(0, 10, txt=result)
                pdf_output = BytesIO()
                pdf.output(pdf_output)
                pdf_output.seek(0)
                st.download_button(label="Download PDF", data=pdf_output, file_name="report.pdf")

            # DOCX 저장
            elif export_format == 'docx':
                doc = docx.Document()
                for idx, result in enumerate(st.session_state['results']):
                    doc.add_heading(f"제목: {st.session_state['df'].at[idx, '제목']}", level=1)
                    doc.add_paragraph(result)
                doc_output = BytesIO()
                doc.save(doc_output)
                doc_output.seek(0)
                st.download_button(label="Download DOCX", data=doc_output, file_name="report.docx")

            # XLSX 저장
            elif export_format == 'xlsx':
                xlsx_output = BytesIO()
                workbook = xlsxwriter.Workbook(xlsx_output, {'in_memory': True})
                worksheet = workbook.add_worksheet()
                worksheet.write(0, 0, "제목")
                worksheet.write(0, 1, "LLM 결과")
                for idx, result in enumerate(st.session_state['results']):
                    worksheet.write(idx + 1, 0, st.session_state['df'].at[idx, '제목'])
                    worksheet.write(idx + 1, 1, result)
                workbook.close()
                xlsx_output.seek(0)
                st.download_button(label="Download XLSX", data=xlsx_output, file_name="report.xlsx")

            # TXT 저장
            elif export_format == 'txt':
                txt_output = ""
                for idx, result in enumerate(st.session_state['results']):
                    txt_output += f"제목: {st.session_state['df'].at[idx, '제목']}\n"
                    txt_output += f"{result}\n\n"
                st.download_button(label="Download TXT", data=txt_output, file_name="report.txt")
